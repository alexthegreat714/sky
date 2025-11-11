#!/usr/bin/env python3
import docx
import sys
import os

def extract_text_from_docx(filepath):
    """Extract all text from a .docx file"""
    try:
        doc = docx.Document(filepath)
        text_content = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)

        return '\n'.join(text_content)
    except Exception as e:
        return f"Error reading {filepath}: {str(e)}"

if __name__ == "__main__":
    if len(sys.argv) > 1:
        filepath = sys.argv[1]
        print(extract_text_from_docx(filepath))
    else:
        print("Usage: python3 extract_docx.py <filepath>")
